import os
import sys
import re
import json
import shutil
from datetime import datetime
from tkinter import Tk
from tkinter.filedialog import askopenfilename

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from openpyxl import load_workbook
from openpyxl.styles import PatternFill


# =====================================================
# CAMINHO DO PROJETO E BANCO DE ESCRITA
# =====================================================

CAMINHO_PROJETO = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..")
)

sys.path.insert(0, CAMINHO_PROJETO)

try:
    from ia.leitor_estilo import gerar_orientacao_estilo
except Exception:
    def gerar_orientacao_estilo():
        return (
            "Utilizar linguagem institucional, pedagógica, clara e adequada para avaliação do MEC, "
            "com foco em monitoramento longitudinal, NDE, PPC, melhoria contínua e gestão acadêmica baseada em evidências."
        )


# =====================================================
# CONFIGURAÇÃO E PASTAS
# =====================================================

def carregar_config():
    caminho = os.path.join(CAMINHO_PROJETO, "config.json")

    if not os.path.exists(caminho):
        raise FileNotFoundError("Arquivo config.json não encontrado na raiz do projeto.")

    with open(caminho, "r", encoding="utf-8") as f:
        return json.load(f)


def criar_pasta_projeto(config):
    agora = datetime.now().strftime("%Y_%m_%d_%Hh%M")
    nome_avaliacao = config.get("nome_avaliacao", "TP").replace(" ", "_")
    pasta_base = os.path.join("projetos_gerados", f"{nome_avaliacao}_{agora}")

    pastas = {
        "base": pasta_base,
        "tabelas": os.path.join(pasta_base, "tabelas"),
        "graficos": os.path.join(pasta_base, "graficos"),
        "relatorios": os.path.join(pasta_base, "relatorios"),
        "dados": os.path.join(pasta_base, "dados_processados"),
        "original": os.path.join(pasta_base, "planilha_original")
    }

    for pasta in pastas.values():
        os.makedirs(pasta, exist_ok=True)

    return pastas


# =====================================================
# SELEÇÃO E LEITURA DE PLANILHAS
# =====================================================

def pedir_planilha(config):
    print(f"\nROBÔ MEC - {config.get('nome_avaliacao', 'Teste de Progresso')}")
    print("Selecione a planilha Excel com os resultados.\n")

    root = Tk()
    root.withdraw()
    root.attributes("-topmost", True)

    caminho = askopenfilename(
        title="Selecione a planilha do Teste de Progresso",
        filetypes=[
            ("Arquivos Excel", "*.xlsx"),
            ("Arquivos Excel antigos", "*.xls"),
            ("Todos os arquivos", "*.*")
        ]
    )

    root.destroy()

    if not caminho:
        raise Exception("Nenhuma planilha foi selecionada.")

    print(f"Planilha selecionada: {caminho}")
    return caminho


def carregar_dados(caminho_planilha):
    xls = pd.ExcelFile(caminho_planilha)
    abas = xls.sheet_names

    print("\nAbas encontradas:")
    for aba in abas:
        print(f"- {aba}")

    return pd.read_excel(caminho_planilha, sheet_name=abas[0])


def detectar_colunas(base):
    coluna_nome = None
    coluna_periodo = None

    for col in base.columns:
        col_lower = str(col).lower()

        if "nome" in col_lower or "aluno" in col_lower or "estudante" in col_lower or "discente" in col_lower:
            coluna_nome = col

        if "período" in col_lower or "periodo" in col_lower or "turma" in col_lower:
            coluna_periodo = col

    colunas_numericas = base.select_dtypes(include=[np.number]).columns.tolist()

    if not colunas_numericas:
        raise ValueError("Não encontrei colunas numéricas com acertos ou notas.")

    coluna_total = colunas_numericas[-1]

    print("\nColunas detectadas:")
    print(f"Aluno: {coluna_nome}")
    print(f"Turma/Período: {coluna_periodo}")
    print(f"Coluna usada para desempenho geral: {coluna_total}")

    return coluna_nome, coluna_periodo, coluna_total


# =====================================================
# PROCESSAMENTO DA BASE
# =====================================================

def preparar_base(base, coluna_nome, coluna_periodo, coluna_total, config):
    base = base.copy()

    if coluna_nome is None:
        base["Aluno"] = [f"Aluno_{i + 1}" for i in range(len(base))]
        coluna_nome = "Aluno"

    if coluna_periodo is None:
        base["Periodo"] = "Geral"
        coluna_periodo = "Periodo"

    numero_questoes = config.get("numero_questoes", 120)

    base["Total_acertos"] = pd.to_numeric(base[coluna_total], errors="coerce")
    base = base.dropna(subset=["Total_acertos"])

    maior_valor = base["Total_acertos"].max()

    if maior_valor <= 1:
        base["Percentual_acertos"] = base["Total_acertos"] * 100
    elif maior_valor <= 100:
        base["Percentual_acertos"] = base["Total_acertos"]
    else:
        base["Percentual_acertos"] = (base["Total_acertos"] / numero_questoes) * 100

    base["Media_turma"] = base.groupby(coluna_periodo)["Percentual_acertos"].transform("mean")
    base["Diferenca_media_turma"] = base["Percentual_acertos"] - base["Media_turma"]

    base["Percentual_abaixo_media_turma"] = np.where(
        base["Percentual_acertos"] < base["Media_turma"],
        ((base["Media_turma"] - base["Percentual_acertos"]) / base["Media_turma"]) * 100,
        0
    )

    base["Percentual_acima_media_turma"] = np.where(
        base["Percentual_acertos"] > base["Media_turma"],
        ((base["Percentual_acertos"] - base["Media_turma"]) / base["Media_turma"]) * 100,
        0
    )

    risco = config.get("criterios_risco", {})
    limite_ped = risco.get("risco_pedagogico_ate_percentual_abaixo_media", 10)
    limite_crit = risco.get("risco_critico_a_partir_percentual_abaixo_media", 11)

    base["Nivel_risco"] = np.where(
        base["Percentual_abaixo_media_turma"] >= limite_crit,
        "Risco crítico",
        np.where(
            (base["Percentual_abaixo_media_turma"] > 0) &
            (base["Percentual_abaixo_media_turma"] <= limite_ped),
            "Risco pedagógico",
            "Sem risco pela média da turma"
        )
    )

    base["Classificacao"] = np.where(
        base["Percentual_acertos"] < 60,
        "Atenção pedagógica",
        "Desempenho adequado"
    )

    return base, coluna_nome, coluna_periodo


# =====================================================
# CICLOS FORMATIVOS E PONTOS EXTRAS
# =====================================================

def extrair_numero_periodo(valor):
    numeros = re.findall(r"\d+", str(valor))
    if not numeros:
        return None
    return numeros[0]


def classificar_ciclo(valor_periodo, config):
    numero = extrair_numero_periodo(valor_periodo)
    ciclos = config.get("ciclos_formativos", {})

    for nome_ciclo, periodos in ciclos.items():
        if numero in [str(p) for p in periodos]:
            return nome_ciclo

    return "Não classificado"


def gerar_tabela_ciclos(base, coluna_periodo, config):
    temp = base.copy()
    temp["Ciclo_formativo"] = temp[coluna_periodo].apply(lambda x: classificar_ciclo(x, config))

    return (
        temp.groupby("Ciclo_formativo")
        .agg(
            Numero_estudantes=("Percentual_acertos", "count"),
            Media_percentual=("Percentual_acertos", "mean"),
            Mediana_percentual=("Percentual_acertos", "median"),
            Desvio_padrao=("Percentual_acertos", "std"),
            Minimo=("Percentual_acertos", "min"),
            Maximo=("Percentual_acertos", "max")
        )
        .reset_index()
        .sort_values("Media_percentual", ascending=True)
    )


def gerar_pontos_extras(base, coluna_nome, coluna_periodo, config):
    pontos = base.copy()

    pontos["Faixa_bonus"] = "Sem bônus"
    pontos["Ponto_extra_por_disciplina"] = 0.0
    pontos["Numero_disciplinas_bonus"] = 0
    pontos["Total_pontos_extras"] = 0.0

    for regra in config.get("criterios_pontos_extras", []):
        minimo = regra.get("min_percentual_acima", 0)
        maximo = regra.get("max_percentual_acima", None)

        if maximo is None:
            mascara = pontos["Percentual_acima_media_turma"] >= minimo
        else:
            mascara = (
                (pontos["Percentual_acima_media_turma"] >= minimo) &
                (pontos["Percentual_acima_media_turma"] <= maximo)
            )

        pontos.loc[mascara, "Faixa_bonus"] = regra.get("faixa", "")
        pontos.loc[mascara, "Ponto_extra_por_disciplina"] = regra.get("ponto_por_disciplina", 0)
        pontos.loc[mascara, "Numero_disciplinas_bonus"] = regra.get("numero_disciplinas", 0)
        pontos.loc[mascara, "Total_pontos_extras"] = regra.get("total_pontos", 0)

    pontos_bonus = pontos[pontos["Faixa_bonus"] != "Sem bônus"].copy()

    colunas = [
        coluna_nome,
        coluna_periodo,
        "Percentual_acertos",
        "Media_turma",
        "Percentual_acima_media_turma",
        "Faixa_bonus",
        "Ponto_extra_por_disciplina",
        "Numero_disciplinas_bonus",
        "Total_pontos_extras"
    ]

    return pontos_bonus[colunas].sort_values(
        ["Total_pontos_extras", "Percentual_acima_media_turma"],
        ascending=False
    )


# =====================================================
# GRANDES ÁREAS
# =====================================================

def identificar_grandes_areas(base, config):
    areas = config.get("grandes_areas", {})
    colunas_areas = {}

    for nome_area, termos in areas.items():
        for coluna in base.columns:
            coluna_lower = str(coluna).lower()

            if any(str(termo).lower() in coluna_lower for termo in termos):
                if pd.api.types.is_numeric_dtype(base[coluna]):
                    colunas_areas[nome_area] = coluna
                break

    return colunas_areas


def gerar_tabela_grandes_areas(base, coluna_periodo, colunas_areas):
    tabelas = {}

    for nome_area, coluna_area in colunas_areas.items():
        temp = base.copy()
        temp[coluna_area] = pd.to_numeric(temp[coluna_area], errors="coerce")

        max_area = temp[coluna_area].max()

        if max_area <= 1:
            temp["Percentual_area"] = temp[coluna_area] * 100
        elif max_area <= 100:
            temp["Percentual_area"] = temp[coluna_area]
        else:
            temp["Percentual_area"] = (temp[coluna_area] / max_area) * 100

        tabela = (
            temp.groupby(coluna_periodo)
            .agg(
                Numero_estudantes=(coluna_area, "count"),
                Media_acertos=(coluna_area, "mean"),
                Percentual_medio=("Percentual_area", "mean"),
                Desvio_padrao=("Percentual_area", "std")
            )
            .reset_index()
            .sort_values("Percentual_medio", ascending=False)
        )

        tabelas[nome_area] = tabela

    return tabelas


# =====================================================
# PLANO DE AÇÃO
# =====================================================

def gerar_plano_acao(ranking_periodos, alunos_risco, tabelas_areas, config):
    plano = []

    if len(alunos_risco) > 0:
        plano.append({
            "Problema identificado": "Estudantes em risco pedagógico ou crítico",
            "Ação proposta": "Implementar devolutivas formativas, monitorias acadêmicas e acompanhamento individualizado.",
            "Responsável": "Coordenação / NDE / Docentes",
            "Prazo": "Próximo ciclo avaliativo",
            "Indicador": "Redução do número de estudantes em risco"
        })

    for nome_area, tabela_area in tabelas_areas.items():
        plano.append({
            "Problema identificado": f"Menor desempenho relativo em {nome_area}",
            "Ação proposta": f"Revisar conteúdos, estratégias de ensino e instrumentos avaliativos vinculados ao eixo de {nome_area}.",
            "Responsável": "NDE / Docentes da área",
            "Prazo": "Próximo semestre",
            "Indicador": f"Elevação do percentual médio em {nome_area}"
        })

    if not ranking_periodos.empty:
        diferenca = ranking_periodos["Media_percentual"].max() - ranking_periodos["Media_percentual"].min()

        if diferenca >= 10:
            plano.append({
                "Problema identificado": "Diferença expressiva de desempenho entre turmas/períodos",
                "Ação proposta": "Discutir a variabilidade no NDE e planejar ações de nivelamento por período.",
                "Responsável": "Coordenação / NDE",
                "Prazo": "Próxima reunião do NDE",
                "Indicador": "Redução da diferença entre médias das turmas"
            })

    if not plano:
        plano.append({
            "Problema identificado": "Manutenção do acompanhamento longitudinal",
            "Ação proposta": "Manter análise semestral dos resultados do Teste de Progresso.",
            "Responsável": "Coordenação / NDE",
            "Prazo": "Contínuo",
            "Indicador": "Série histórica de desempenho por período"
        })

    return pd.DataFrame(plano)


# =====================================================
# EXCEL
# =====================================================

def aplicar_cores_excel(caminho_excel):
    wb = load_workbook(caminho_excel)

    amarelo = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")
    vermelho = PatternFill(start_color="F4CCCC", end_color="F4CCCC", fill_type="solid")
    verde = PatternFill(start_color="D9EAD3", end_color="D9EAD3", fill_type="solid")

    for aba in wb.sheetnames:
        ws = wb[aba]
        cabecalhos = [cell.value for cell in ws[1]]

        if "Nivel_risco" in cabecalhos:
            coluna_risco = cabecalhos.index("Nivel_risco") + 1

            for linha in range(2, ws.max_row + 1):
                valor = ws.cell(row=linha, column=coluna_risco).value

                if valor == "Risco pedagógico":
                    for col in range(1, ws.max_column + 1):
                        ws.cell(row=linha, column=col).fill = amarelo

                elif valor == "Risco crítico":
                    for col in range(1, ws.max_column + 1):
                        ws.cell(row=linha, column=col).fill = vermelho

        if aba == "Pontos_extras":
            for linha in range(2, ws.max_row + 1):
                for col in range(1, ws.max_column + 1):
                    ws.cell(row=linha, column=col).fill = verde

    wb.save(caminho_excel)


def gerar_tabelas(base, coluna_nome, coluna_periodo, pastas, config):
    ranking_alunos = base.sort_values("Percentual_acertos", ascending=False)

    ranking_periodos = (
        base.groupby(coluna_periodo)
        .agg(
            N_estudantes=(coluna_nome, "count"),
            Media_percentual=("Percentual_acertos", "mean"),
            Mediana_percentual=("Percentual_acertos", "median"),
            Desvio_padrao=("Percentual_acertos", "std"),
            Minimo=("Percentual_acertos", "min"),
            Maximo=("Percentual_acertos", "max")
        )
        .reset_index()
        .sort_values("Media_percentual", ascending=False)
    )

    tabela_ciclos = gerar_tabela_ciclos(base, coluna_periodo, config)

    colunas_areas = identificar_grandes_areas(base, config)
    tabelas_areas = gerar_tabela_grandes_areas(base, coluna_periodo, colunas_areas)

    alunos_risco = base[
        base["Nivel_risco"].isin(["Risco pedagógico", "Risco crítico"])
    ].copy()

    pontos_extras = gerar_pontos_extras(base, coluna_nome, coluna_periodo, config)

    plano_acao = gerar_plano_acao(
        ranking_periodos,
        alunos_risco,
        tabelas_areas,
        config
    )

    caminho_excel = os.path.join(pastas["tabelas"], "tabelas_resultados_mec.xlsx")

    with pd.ExcelWriter(caminho_excel, engine="openpyxl") as writer:
        base.to_excel(writer, sheet_name="Base_processada", index=False)
        ranking_alunos.to_excel(writer, sheet_name="Ranking_alunos", index=False)
        ranking_periodos.to_excel(writer, sheet_name="Ranking_periodos", index=False)
        tabela_ciclos.to_excel(writer, sheet_name="Ciclos_formativos", index=False)
        alunos_risco.to_excel(writer, sheet_name="Alunos_em_risco", index=False)
        pontos_extras.to_excel(writer, sheet_name="Pontos_extras", index=False)
        plano_acao.to_excel(writer, sheet_name="Plano_acao", index=False)

        for nome_area, tabela_area in tabelas_areas.items():
            tabela_area.to_excel(writer, sheet_name=nome_area[:28], index=False)

    aplicar_cores_excel(caminho_excel)

    base.to_csv(
        os.path.join(pastas["dados"], "base_processada.csv"),
        sep=";",
        decimal=",",
        index=False,
        encoding="utf-8-sig"
    )

    pontos_extras.to_csv(
        os.path.join(pastas["dados"], "pontos_extras.csv"),
        sep=";",
        decimal=",",
        index=False,
        encoding="utf-8-sig"
    )

    plano_acao.to_csv(
        os.path.join(pastas["dados"], "plano_acao.csv"),
        sep=";",
        decimal=",",
        index=False,
        encoding="utf-8-sig"
    )

    return (
        ranking_alunos,
        ranking_periodos,
        tabela_ciclos,
        alunos_risco,
        pontos_extras,
        plano_acao,
        caminho_excel,
        tabelas_areas
    )


# =====================================================
# GRÁFICOS
# =====================================================

def configurar_estilo_grafico():
    plt.rcParams["font.family"] = "Arial"
    plt.rcParams["axes.spines.top"] = False
    plt.rcParams["axes.spines.right"] = False


def salvar_grafico(nome, pastas):
    png = os.path.join(pastas["graficos"], f"{nome}.png")
    svg = os.path.join(pastas["graficos"], f"{nome}.svg")

    plt.savefig(png, dpi=300, bbox_inches="tight")
    plt.savefig(svg, bbox_inches="tight")
    plt.close()

    return png


def nome_arquivo_seguro(texto):
    texto = texto.lower()
    texto = texto.replace("á", "a").replace("ã", "a").replace("â", "a").replace("à", "a")
    texto = texto.replace("é", "e").replace("ê", "e")
    texto = texto.replace("í", "i")
    texto = texto.replace("ó", "o").replace("õ", "o").replace("ô", "o")
    texto = texto.replace("ú", "u")
    texto = texto.replace("ç", "c")
    texto = re.sub(r"[^a-z0-9]+", "_", texto)
    return texto.strip("_")


def gerar_graficos(base, ranking_periodos, tabela_ciclos, coluna_periodo, pastas):
    configurar_estilo_grafico()

    graficos = {}

    tabela = ranking_periodos.sort_values("Media_percentual", ascending=True)

    plt.figure(figsize=(12, 7))
    barras = plt.bar(
        tabela[coluna_periodo].astype(str),
        tabela["Media_percentual"],
        color="#28a9d6"
    )

    plt.plot(
        tabela[coluna_periodo].astype(str),
        tabela["Media_percentual"],
        color="#222222",
        marker="o",
        linewidth=3
    )

    for barra in barras:
        valor = barra.get_height()
        plt.text(
            barra.get_x() + barra.get_width() / 2,
            valor + 1,
            f"{valor:.1f}%",
            ha="center",
            fontsize=12,
            fontweight="bold"
        )

    plt.title("Desempenho médio por turma/período", fontsize=18, fontweight="bold")
    plt.xlabel("Turma/Período")
    plt.ylabel("Média de acertos (%)")
    plt.grid(axis="y", alpha=0.2)
    plt.ylim(0, max(tabela["Media_percentual"]) + 15)
    graficos["ranking"] = salvar_grafico("ranking_periodos_visual", pastas)

    plt.figure(figsize=(11, 6))
    plt.hist(
        base["Percentual_acertos"],
        bins=15,
        color="#28a9d6",
        edgecolor="white"
    )

    media = base["Percentual_acertos"].mean()
    plt.axvline(media, color="#ec174c", linestyle="--", linewidth=3)

    plt.title("Distribuição geral do desempenho dos estudantes", fontsize=18, fontweight="bold")
    plt.xlabel("Percentual de acertos")
    plt.ylabel("Número de estudantes")
    plt.grid(axis="y", alpha=0.2)
    graficos["histograma"] = salvar_grafico("histograma_desempenho_visual", pastas)

    plt.figure(figsize=(12, 6))
    base.boxplot(column="Percentual_acertos", by=coluna_periodo, grid=False)
    plt.title("Comparação do desempenho por turma/período", fontsize=18, fontweight="bold")
    plt.suptitle("")
    plt.xlabel("Turma/Período")
    plt.ylabel("Percentual de acertos")
    plt.grid(axis="y", alpha=0.2)
    graficos["boxplot"] = salvar_grafico("boxplot_periodos_visual", pastas)

    if not tabela_ciclos.empty:
        ciclo_plot = tabela_ciclos.sort_values("Media_percentual", ascending=True)

        plt.figure(figsize=(10, 6))
        barras = plt.bar(
            ciclo_plot["Ciclo_formativo"].astype(str),
            ciclo_plot["Media_percentual"],
            color="#36d6b0"
        )

        for barra in barras:
            valor = barra.get_height()
            plt.text(
                barra.get_x() + barra.get_width() / 2,
                valor + 1,
                f"{valor:.1f}%",
                ha="center",
                fontsize=12,
                fontweight="bold"
            )

        plt.title("Desempenho médio por ciclo formativo", fontsize=18, fontweight="bold")
        plt.xlabel("Ciclo formativo")
        plt.ylabel("Média de acertos (%)")
        plt.grid(axis="y", alpha=0.2)
        plt.ylim(0, max(ciclo_plot["Media_percentual"]) + 15)
        graficos["ciclos"] = salvar_grafico("ciclos_formativos_visual", pastas)

    return graficos


def gerar_grafico_area(nome_area, tabela_area, coluna_periodo, pastas):
    configurar_estilo_grafico()

    nome_base = nome_arquivo_seguro(nome_area)
    tabela = tabela_area.sort_values("Percentual_medio", ascending=True)

    plt.figure(figsize=(12, 7))
    barras = plt.bar(
        tabela[coluna_periodo].astype(str),
        tabela["Percentual_medio"],
        color="#28a9d6"
    )

    plt.plot(
        tabela[coluna_periodo].astype(str),
        tabela["Percentual_medio"],
        color="#222222",
        marker="o",
        linewidth=3
    )

    for barra in barras:
        valor = barra.get_height()
        plt.text(
            barra.get_x() + barra.get_width() / 2,
            valor + 1,
            f"{valor:.1f}%",
            ha="center",
            fontsize=12,
            fontweight="bold"
        )

    plt.title(f"Percentual médio em {nome_area} por turma/período", fontsize=16, fontweight="bold")
    plt.xlabel("Turma/Período")
    plt.ylabel("Percentual médio (%)")
    plt.grid(axis="y", alpha=0.2)
    plt.ylim(0, max(tabela["Percentual_medio"]) + 15)

    caminho_percentual = os.path.join(pastas["graficos"], f"area_{nome_base}_percentual.png")
    plt.savefig(caminho_percentual, dpi=300, bbox_inches="tight")
    plt.close()

    plt.figure(figsize=(12, 7))
    barras = plt.bar(
        tabela[coluna_periodo].astype(str),
        tabela["Media_acertos"],
        color="#36d6b0"
    )

    for barra in barras:
        valor = barra.get_height()
        plt.text(
            barra.get_x() + barra.get_width() / 2,
            valor + 0.5,
            f"{valor:.1f}",
            ha="center",
            fontsize=12,
            fontweight="bold"
        )

    plt.title(f"Média de acertos em {nome_area} por turma/período", fontsize=16, fontweight="bold")
    plt.xlabel("Turma/Período")
    plt.ylabel("Média de acertos")
    plt.grid(axis="y", alpha=0.2)

    caminho_media = os.path.join(pastas["graficos"], f"area_{nome_base}_media_acertos.png")
    plt.savefig(caminho_media, dpi=300, bbox_inches="tight")
    plt.close()

    return caminho_percentual, caminho_media


# =====================================================
# FORMATAÇÃO PROFISSIONAL DO WORD
# =====================================================

def configurar_documento(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)


def add_titulo(doc, numero, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{numero}. {texto.upper()}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    return p


def add_subtitulo(doc, numero, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{numero} {texto}")
    run.bold = True
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def add_paragrafo(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(texto)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")

    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge_name, edge_data in {
        "top": top,
        "bottom": bottom,
        "left": left,
        "right": right,
    }.items():
        tag = "w:{}".format(edge_name)
        element = borders.find(qn(tag))

        if element is None:
            element = OxmlElement(tag)
            borders.append(element)

        if edge_data:
            element.set(qn("w:val"), edge_data.get("val", "single"))
            element.set(qn("w:sz"), str(edge_data.get("sz", 6)))
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), edge_data.get("color", "000000"))
        else:
            element.set(qn("w:val"), "nil")


def formatar_tabela_ibge(tabela):
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.autofit = True

    n_linhas = len(tabela.rows)

    for i, row in enumerate(tabela.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)

            set_cell_border(cell, left=None, right=None)

            if i == 0:
                set_cell_border(
                    cell,
                    top={"val": "single", "sz": 8, "color": "000000"},
                    bottom={"val": "single", "sz": 8, "color": "000000"},
                    left=None,
                    right=None,
                )

                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            elif i == n_linhas - 1:
                set_cell_border(
                    cell,
                    top=None,
                    bottom={"val": "single", "sz": 8, "color": "000000"},
                    left=None,
                    right=None,
                )

            else:
                set_cell_border(cell, top=None, bottom=None, left=None, right=None)

            if i > 0 and i % 2 == 1:
                set_cell_shading(cell, "F2F2F2")
            else:
                set_cell_shading(cell, "FFFFFF")


def adicionar_titulo_tabela(doc, numero, titulo):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Tabela {numero} – {titulo}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def adicionar_fonte(doc, config):
    fonte = config.get("relatorio", {}).get(
        "fonte_tabelas_figuras",
        "Fonte: Elaboração própria a partir dos resultados do Teste de Progresso."
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(fonte)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def adicionar_legenda_figura(doc, numero, texto, config):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Figura {numero} – {texto}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    adicionar_fonte(doc, config)


# =====================================================
# RELATÓRIO WORD
# =====================================================

def criar_relatorio_word(
    base,
    ranking_periodos,
    tabela_ciclos,
    alunos_risco,
    pontos_extras,
    plano_acao,
    graficos,
    coluna_nome,
    coluna_periodo,
    pastas,
    tabelas_areas,
    config
):
    total = len(base)
    media_geral = base["Percentual_acertos"].mean()

    melhor = ranking_periodos.iloc[0]
    pior = ranking_periodos.iloc[-1]

    nome_instituicao = config.get("nome_instituicao", "")
    nome_curso = config.get("nome_curso", "")
    nome_avaliacao = config.get("nome_avaliacao", "Teste de Progresso")
    titulo = config.get("relatorio", {}).get("titulo", "Relatório Pedagógico do Teste de Progresso")
    finalidade = config.get("relatorio", {}).get("finalidade", "")
    estilo = config.get("estilo_mec", {})

    orientacao_estilo = gerar_orientacao_estilo()

    doc = Document()
    configurar_documento(doc)

    doc.add_heading(titulo, level=0)

    add_paragrafo(doc, f"Instituição: {nome_instituicao}")
    add_paragrafo(doc, f"Curso: {nome_curso}")
    add_paragrafo(doc, f"Instrumento avaliativo: {nome_avaliacao}")

    add_paragrafo(
        doc,
        f"O {nome_avaliacao} é utilizado como ferramenta de avaliação cognitiva longitudinal, "
        f"permitindo acompanhar a progressão do estudante ao longo da formação médica. "
        f"Neste relatório, os resultados são apresentados de forma didática e analítica, com foco "
        f"na evolução entre períodos, nos ciclos formativos, nas grandes áreas do conhecimento, "
        f"nos estudantes que demandam acompanhamento pedagógico, nos estudantes elegíveis à bonificação "
        f"acadêmica e no plano de ação pedagógico decorrente dos indicadores. {finalidade}"
    )

    add_paragrafo(doc, orientacao_estilo)

    add_titulo(doc, "1", "Resumo executivo")

    add_paragrafo(
        doc,
        f"O presente relatório sintetiza os resultados do {nome_avaliacao} aplicado aos estudantes do curso "
        f"de {nome_curso} da {nome_instituicao}. Foram avaliados {total} estudantes, com média geral "
        f"de {media_geral:.2f}% de acertos."
    )

    add_paragrafo(
        doc,
        f"O maior desempenho médio foi identificado em {melhor[coluna_periodo]}, com média de "
        f"{melhor['Media_percentual']:.2f}%, enquanto o menor desempenho médio foi observado em "
        f"{pior[coluna_periodo]}, com média de {pior['Media_percentual']:.2f}%."
    )

    add_paragrafo(
        doc,
        f"Foram identificados {len(alunos_risco)} estudantes classificados em risco pedagógico ou crítico. "
        f"Também foram identificados {len(pontos_extras)} estudantes elegíveis à bonificação acadêmica por "
        f"desempenho acima da média da própria turma."
    )

    add_paragrafo(
        doc,
        "Os indicadores produzidos subsidiam ações da coordenação, do Núcleo Docente Estruturante (NDE) "
        "e do colegiado, com foco no acompanhamento longitudinal, na melhoria contínua e no fortalecimento "
        "da gestão acadêmica baseada em evidências."
    )

    add_titulo(doc, "2", "Leitura pedagógica geral dos resultados")

    add_paragrafo(
        doc,
        f"Foram avaliados {total} estudantes. Conforme apresentado na Tabela 1 e nas Figuras 1, 2 e 3, "
        f"o maior desempenho médio foi observado em {melhor[coluna_periodo]}, com média de "
        f"{melhor['Media_percentual']:.2f}% de acertos. O menor desempenho médio foi observado em "
        f"{pior[coluna_periodo]}, com média de {pior['Media_percentual']:.2f}% de acertos."
    )

    add_paragrafo(
        doc,
        estilo.get(
            "progressao",
            "Observa-se progressão longitudinal do desempenho médio ao longo dos períodos do curso."
        )
    )

    add_paragrafo(
        doc,
        estilo.get(
            "competencias",
            "Os resultados sugerem consolidação progressiva das competências cognitivas esperadas para cada etapa formativa."
        )
    )

    add_titulo(doc, "3", "Distribuição geral do desempenho")

    add_paragrafo(
        doc,
        "A distribuição geral do desempenho dos estudantes está apresentada na Figura 1. Esse gráfico permite "
        "visualizar a concentração dos estudantes em torno da média geral, bem como a presença de grupos com "
        "desempenho acima e abaixo do esperado para suas respectivas turmas."
    )

    doc.add_picture(graficos["histograma"], width=Inches(6))
    adicionar_legenda_figura(doc, 1, "Distribuição geral do percentual de acertos dos estudantes.", config)

    add_titulo(doc, "4", "Progressão por turma/período")

    add_paragrafo(
        doc,
        "A progressão do desempenho por turma/período está apresentada na Figura 2 e na Tabela 1. "
        "A Figura 3 complementa essa análise ao demonstrar a dispersão interna dos resultados em cada grupo. "
        "Em uma avaliação longitudinal, espera-se que os estudantes apresentem aumento gradual das taxas de acerto "
        "conforme avançam no curso."
    )

    doc.add_picture(graficos["ranking"], width=Inches(6))
    adicionar_legenda_figura(doc, 2, "Ranking dos períodos/turmas segundo a média percentual de acertos.", config)

    doc.add_picture(graficos["boxplot"], width=Inches(6))
    adicionar_legenda_figura(doc, 3, "Distribuição do desempenho por turma/período.", config)

    adicionar_titulo_tabela(doc, 1, "Síntese estatística do desempenho por turma/período")

    tabela_doc = doc.add_table(rows=1, cols=6)
    cab = tabela_doc.rows[0].cells
    cab[0].text = "Turma/Período"
    cab[1].text = "N"
    cab[2].text = "Média (%)"
    cab[3].text = "Mediana (%)"
    cab[4].text = "Desvio-padrão"
    cab[5].text = "Máximo (%)"

    for _, row in ranking_periodos.iterrows():
        cells = tabela_doc.add_row().cells
        cells[0].text = str(row[coluna_periodo])
        cells[1].text = str(int(row["N_estudantes"]))
        cells[2].text = f"{row['Media_percentual']:.2f}"
        cells[3].text = f"{row['Mediana_percentual']:.2f}"
        cells[4].text = f"{row['Desvio_padrao']:.2f}"
        cells[5].text = f"{row['Maximo']:.2f}"

    formatar_tabela_ibge(tabela_doc)
    adicionar_fonte(doc, config)

    add_titulo(doc, "5", "Resultados por ciclos formativos")

    add_paragrafo(
        doc,
        "A análise por ciclos formativos está apresentada na Figura 4 e na Tabela 2. Essa organização permite "
        "verificar se o desempenho médio acompanha a lógica esperada de progressão formativa entre Ciclo Básico, "
        "Ciclo Clínico e Internato."
    )

    if "ciclos" in graficos:
        doc.add_picture(graficos["ciclos"], width=Inches(6))
        adicionar_legenda_figura(doc, 4, "Desempenho médio por ciclo formativo.", config)

    adicionar_titulo_tabela(doc, 2, "Síntese do desempenho por ciclo formativo")

    tabela_ciclo_doc = doc.add_table(rows=1, cols=6)
    cab = tabela_ciclo_doc.rows[0].cells
    cab[0].text = "Ciclo"
    cab[1].text = "N"
    cab[2].text = "Média (%)"
    cab[3].text = "Mediana (%)"
    cab[4].text = "Desvio-padrão"
    cab[5].text = "Máximo (%)"

    for _, row in tabela_ciclos.iterrows():
        cells = tabela_ciclo_doc.add_row().cells
        cells[0].text = str(row["Ciclo_formativo"])
        cells[1].text = str(int(row["Numero_estudantes"]))
        cells[2].text = f"{row['Media_percentual']:.2f}"
        cells[3].text = f"{row['Mediana_percentual']:.2f}"
        cells[4].text = f"{row['Desvio_padrao']:.2f}"
        cells[5].text = f"{row['Maximo']:.2f}"

    formatar_tabela_ibge(tabela_ciclo_doc)
    adicionar_fonte(doc, config)

    add_titulo(doc, "6", "Resultados por grandes áreas do conhecimento")

    add_paragrafo(
        doc,
        "Os resultados por grandes áreas do conhecimento estão apresentados nas tabelas e figuras desta seção. "
        "Cada eixo foi analisado considerando número de estudantes, média de acertos, percentual médio e desvio-padrão."
    )

    numero_tabela = 3
    numero_figura = 5

    if not tabelas_areas:
        add_paragrafo(doc, "Não foram identificadas automaticamente colunas referentes às grandes áreas do conhecimento.")
    else:
        for idx, (nome_area, tabela_area) in enumerate(tabelas_areas.items(), start=1):
            add_subtitulo(doc, f"6.{idx}", nome_area)

            grafico_percentual, grafico_media = gerar_grafico_area(
                nome_area,
                tabela_area,
                coluna_periodo,
                pastas
            )

            melhor_area = tabela_area.sort_values("Percentual_medio", ascending=False).iloc[0]
            menor_area = tabela_area.sort_values("Percentual_medio", ascending=True).iloc[0]

            add_paragrafo(
                doc,
                f"Os resultados referentes à área de {nome_area} estão apresentados nas Figuras "
                f"{numero_figura} e {numero_figura + 1} e na Tabela {numero_tabela}. "
                f"No eixo de {nome_area}, a taxa média de acerto variou de "
                f"{menor_area['Percentual_medio']:.2f}% a {melhor_area['Percentual_medio']:.2f}% "
                f"entre as turmas/períodos avaliados."
            )

            add_paragrafo(
                doc,
                f"O maior desempenho foi observado em {melhor_area[coluna_periodo]}, com média de "
                f"{melhor_area['Media_acertos']:.2f} acertos. O menor desempenho foi observado em "
                f"{menor_area[coluna_periodo]}, com média de {menor_area['Media_acertos']:.2f} acertos. "
                f"Esses dados subsidiam a análise do alcance das competências relacionadas ao eixo de {nome_area}."
            )

            doc.add_picture(grafico_percentual, width=Inches(6))
            adicionar_legenda_figura(doc, numero_figura, f"Percentual médio em {nome_area} por turma/período.", config)
            numero_figura += 1

            doc.add_picture(grafico_media, width=Inches(6))
            adicionar_legenda_figura(doc, numero_figura, f"Média de acertos em {nome_area} por turma/período.", config)
            numero_figura += 1

            adicionar_titulo_tabela(
                doc,
                numero_tabela,
                f"Número de estudantes, média de acertos e percentual médio em {nome_area}"
            )

            tabela_area_doc = doc.add_table(rows=1, cols=5)

            cab = tabela_area_doc.rows[0].cells
            cab[0].text = "Turma/Período"
            cab[1].text = "Número"
            cab[2].text = "Média de acertos"
            cab[3].text = "Percentual médio (%)"
            cab[4].text = "DP (%)"

            for _, row in tabela_area.sort_values("Percentual_medio", ascending=False).iterrows():
                cells = tabela_area_doc.add_row().cells
                cells[0].text = str(row[coluna_periodo])
                cells[1].text = str(int(row["Numero_estudantes"]))
                cells[2].text = f"{row['Media_acertos']:.2f}"
                cells[3].text = f"{row['Percentual_medio']:.2f}"
                cells[4].text = f"{row['Desvio_padrao']:.2f}"

            formatar_tabela_ibge(tabela_area_doc)
            adicionar_fonte(doc, config)
            numero_tabela += 1

    add_titulo(doc, "7", "Estudantes em risco pedagógico")

    risco_pedagogico = alunos_risco[alunos_risco["Nivel_risco"] == "Risco pedagógico"]
    risco_critico = alunos_risco[alunos_risco["Nivel_risco"] == "Risco crítico"]

    add_paragrafo(
        doc,
        f"A classificação dos estudantes em risco está apresentada na Tabela {numero_tabela}. "
        f"Foram identificados {len(risco_pedagogico)} estudantes com desempenho até 10% abaixo da média "
        f"da própria turma e {len(risco_critico)} estudantes com desempenho a partir de 11% abaixo da média "
        f"da própria turma."
    )

    adicionar_titulo_tabela(doc, numero_tabela, "Estudantes classificados em risco pedagógico ou crítico")

    tabela_risco = doc.add_table(rows=1, cols=5)

    cab = tabela_risco.rows[0].cells
    cab[0].text = "Aluno"
    cab[1].text = "Turma/Período"
    cab[2].text = "% de acertos"
    cab[3].text = "% abaixo da média da turma"
    cab[4].text = "Classificação"

    for _, row in alunos_risco.sort_values("Percentual_abaixo_media_turma", ascending=False).iterrows():
        cells = tabela_risco.add_row().cells
        cells[0].text = str(row[coluna_nome])
        cells[1].text = str(row[coluna_periodo])
        cells[2].text = f"{row['Percentual_acertos']:.2f}%"
        cells[3].text = f"{row['Percentual_abaixo_media_turma']:.2f}%"
        cells[4].text = str(row["Nivel_risco"])

    formatar_tabela_ibge(tabela_risco)
    adicionar_fonte(doc, config)
    numero_tabela += 1

    add_titulo(doc, "8", "Bonificação acadêmica por desempenho acima da média da turma")

    add_paragrafo(
        doc,
        f"A relação de estudantes elegíveis à bonificação acadêmica está apresentada na Tabela {numero_tabela}. "
        f"A bonificação foi calculada a partir da comparação entre o desempenho individual e a média da própria turma."
    )

    adicionar_titulo_tabela(doc, numero_tabela, "Estudantes elegíveis à bonificação acadêmica")

    tabela_bonus = doc.add_table(rows=1, cols=7)

    cab = tabela_bonus.rows[0].cells
    cab[0].text = "Aluno"
    cab[1].text = "Turma/Período"
    cab[2].text = "% acertos"
    cab[3].text = "% acima da média"
    cab[4].text = "Faixa"
    cab[5].text = "Disciplinas"
    cab[6].text = "Total pontos"

    for _, row in pontos_extras.iterrows():
        cells = tabela_bonus.add_row().cells
        cells[0].text = str(row[coluna_nome])
        cells[1].text = str(row[coluna_periodo])
        cells[2].text = f"{row['Percentual_acertos']:.2f}%"
        cells[3].text = f"{row['Percentual_acima_media_turma']:.2f}%"
        cells[4].text = str(row["Faixa_bonus"])
        cells[5].text = str(int(row["Numero_disciplinas_bonus"]))
        cells[6].text = f"{row['Total_pontos_extras']:.2f}"

    formatar_tabela_ibge(tabela_bonus)
    adicionar_fonte(doc, config)
    numero_tabela += 1

    add_titulo(doc, "9", "Plano de ação pedagógico")

    add_paragrafo(
        doc,
        f"O plano de ação pedagógico está apresentado na Tabela {numero_tabela}. Ele foi estruturado a partir "
        f"dos indicadores identificados no Teste de Progresso, considerando estudantes em risco, áreas com menor "
        f"desempenho relativo e diferenças entre turmas/períodos."
    )

    adicionar_titulo_tabela(doc, numero_tabela, "Plano de ação pedagógico baseado nos indicadores do Teste de Progresso")

    tabela_plano = doc.add_table(rows=1, cols=5)

    cab = tabela_plano.rows[0].cells
    cab[0].text = "Problema identificado"
    cab[1].text = "Ação proposta"
    cab[2].text = "Responsável"
    cab[3].text = "Prazo"
    cab[4].text = "Indicador"

    for _, row in plano_acao.iterrows():
        cells = tabela_plano.add_row().cells
        cells[0].text = str(row["Problema identificado"])
        cells[1].text = str(row["Ação proposta"])
        cells[2].text = str(row["Responsável"])
        cells[3].text = str(row["Prazo"])
        cells[4].text = str(row["Indicador"])

    formatar_tabela_ibge(tabela_plano)
    adicionar_fonte(doc, config)

    add_titulo(doc, "10", "Encaminhamentos pedagógicos sugeridos")

    add_paragrafo(
        doc,
        estilo.get(
            "monitoramento",
            "Os indicadores produzidos pelo Teste de Progresso subsidiam ações do NDE e da coordenação do curso."
        )
    )

    add_paragrafo(
        doc,
        estilo.get(
            "reforco",
            "As áreas com menor desempenho são analisadas como oportunidades de reforço pedagógico."
        )
    )

    add_paragrafo(
        doc,
        estilo.get(
            "curriculo",
            "Os resultados observados reforçam o alinhamento entre a matriz curricular e as competências previstas no PPC."
        )
    )

    add_paragrafo(
        doc,
        "Recomenda-se que os resultados sejam discutidos em reuniões do NDE e do colegiado, com registro em ata, "
        "definição de ações de melhoria, acompanhamento dos estudantes em risco e planejamento de intervenções "
        "pedagógicas por turma, ciclo formativo e área do conhecimento."
    )

    caminho = os.path.join(pastas["relatorios"], "relatorio_completo_mec.docx")
    doc.save(caminho)

    return caminho


# =====================================================
# COMPARAÇÃO LONGITUDINAL
# =====================================================

def perguntar_comparacao_longitudinal():
    print("\n========================================")
    print("COMPARAÇÃO LONGITUDINAL")
    print("========================================")
    print("Deseja comparar este Teste de Progresso com outro resultado?")
    print("1 - Não")
    print("2 - Sim, comparar com outro semestre")
    print("3 - Sim, comparar com vários semestres")

    opcao = input("\nEscolha uma opção: ").strip()

    if opcao not in ["2", "3"]:
        return False, []

    quantidade = 1

    if opcao == "3":
        quantidade = int(input("\nQuantas planilhas deseja comparar além da atual? ").strip())

    planilhas = []

    for i in range(quantidade):
        print(f"\nSelecione a planilha de comparação {i + 1}")

        root = Tk()
        root.withdraw()
        root.attributes("-topmost", True)

        caminho = askopenfilename(
            title=f"Selecione a planilha de comparação {i + 1}",
            filetypes=[
                ("Arquivos Excel", "*.xlsx"),
                ("Arquivos Excel antigos", "*.xls"),
                ("Todos os arquivos", "*.*")
            ]
        )

        root.destroy()

        if caminho:
            planilhas.append(caminho)

    return True, planilhas


def executar_comparacao_longitudinal(
    base_atual,
    alunos_risco_atual,
    pontos_extras_atual,
    planilhas_comparacao,
    pastas,
    config
):
    registros = []

    registros.append({
        "Aplicação": "Teste atual",
        "Média geral (%)": base_atual["Percentual_acertos"].mean(),
        "N estudantes": len(base_atual),
        "Alunos em risco": len(alunos_risco_atual),
        "Elegíveis a pontos extras": len(pontos_extras_atual)
    })

    for caminho in planilhas_comparacao:
        try:
            nome = os.path.basename(caminho)
            nome = nome.replace(".xlsx", "").replace(".xls", "")

            base_extra = carregar_dados(caminho)
            coluna_nome_extra, coluna_periodo_extra, coluna_total_extra = detectar_colunas(base_extra)

            base_extra, coluna_nome_extra, coluna_periodo_extra = preparar_base(
                base_extra,
                coluna_nome_extra,
                coluna_periodo_extra,
                coluna_total_extra,
                config
            )

            alunos_risco_extra = base_extra[
                base_extra["Nivel_risco"].isin(["Risco pedagógico", "Risco crítico"])
            ].copy()

            pontos_extras_extra = gerar_pontos_extras(
                base_extra,
                coluna_nome_extra,
                coluna_periodo_extra,
                config
            )

            registros.append({
                "Aplicação": nome,
                "Média geral (%)": base_extra["Percentual_acertos"].mean(),
                "N estudantes": len(base_extra),
                "Alunos em risco": len(alunos_risco_extra),
                "Elegíveis a pontos extras": len(pontos_extras_extra)
            })

        except Exception as erro:
            print(f"Erro ao processar comparação: {caminho}")
            print(erro)

    tabela = pd.DataFrame(registros)

    caminho_excel = os.path.join(
        pastas["tabelas"],
        "comparacao_longitudinal.xlsx"
    )

    tabela.to_excel(caminho_excel, index=False)

    plt.figure(figsize=(12, 6))
    plt.plot(
        tabela["Aplicação"],
        tabela["Média geral (%)"],
        marker="o",
        linewidth=3
    )

    for i, row in tabela.iterrows():
        plt.text(
            i,
            row["Média geral (%)"] + 0.5,
            f"{row['Média geral (%)']:.1f}%",
            ha="center",
            fontsize=11,
            fontweight="bold"
        )

    plt.title("Evolução longitudinal do desempenho médio", fontsize=18, fontweight="bold")
    plt.xlabel("Aplicação do Teste de Progresso")
    plt.ylabel("Média geral (%)")
    plt.grid(alpha=0.2)

    caminho_grafico = os.path.join(
        pastas["graficos"],
        "evolucao_longitudinal.png"
    )

    plt.savefig(caminho_grafico, dpi=300, bbox_inches="tight")
    plt.close()

    caminho_word = os.path.join(
        pastas["relatorios"],
        "relatorio_comparacao_longitudinal.docx"
    )

    doc = Document()
    configurar_documento(doc)

    doc.add_heading("Relatório de Comparação Longitudinal do Teste de Progresso", level=0)

    add_titulo(doc, "1", "Resumo da comparação longitudinal")

    add_paragrafo(
        doc,
        "A comparação longitudinal foi realizada a partir de diferentes aplicações do Teste de Progresso. "
        "Essa análise permite acompanhar a evolução institucional do desempenho discente ao longo do tempo "
        "e subsidia decisões da coordenação, do NDE e do colegiado."
    )

    add_paragrafo(
        doc,
        "Os resultados comparativos estão apresentados na Figura 1 e na Tabela 1. A figura demonstra a evolução "
        "da média geral entre as aplicações analisadas, enquanto a tabela sintetiza o número de estudantes, "
        "os estudantes em risco e os estudantes elegíveis à bonificação acadêmica."
    )

    doc.add_picture(caminho_grafico, width=Inches(6))

    adicionar_legenda_figura(
        doc,
        1,
        "Evolução longitudinal do desempenho médio entre aplicações do Teste de Progresso.",
        config
    )

    adicionar_titulo_tabela(
        doc,
        1,
        "Síntese comparativa entre aplicações do Teste de Progresso"
    )

    tabela_doc = doc.add_table(rows=1, cols=5)

    cab = tabela_doc.rows[0].cells
    cab[0].text = "Aplicação"
    cab[1].text = "Média geral (%)"
    cab[2].text = "N estudantes"
    cab[3].text = "Alunos em risco"
    cab[4].text = "Elegíveis a pontos extras"

    for _, row in tabela.iterrows():
        cells = tabela_doc.add_row().cells
        cells[0].text = str(row["Aplicação"])
        cells[1].text = f"{row['Média geral (%)']:.2f}"
        cells[2].text = str(int(row["N estudantes"]))
        cells[3].text = str(int(row["Alunos em risco"]))
        cells[4].text = str(int(row["Elegíveis a pontos extras"]))

    formatar_tabela_ibge(tabela_doc)
    adicionar_fonte(doc, config)

    add_titulo(doc, "2", "Leitura pedagógica da evolução")

    if len(tabela) >= 2:
        primeira = tabela.iloc[0]["Média geral (%)"]
        ultima = tabela.iloc[-1]["Média geral (%)"]
        diferenca = ultima - primeira

        if diferenca > 0:
            texto = (
                f"Observou-se aumento de {diferenca:.2f} pontos percentuais entre a primeira e a última "
                f"aplicação analisada. Esse resultado sugere evolução do desempenho médio institucional "
                f"no período comparado."
            )
        elif diferenca < 0:
            texto = (
                f"Observou-se redução de {abs(diferenca):.2f} pontos percentuais entre a primeira e a última "
                f"aplicação analisada. Esse resultado deve ser analisado pela coordenação e pelo NDE como "
                f"oportunidade de revisão pedagógica."
            )
        else:
            texto = (
                "As médias permaneceram estáveis entre a primeira e a última aplicação analisada, indicando "
                "manutenção do desempenho médio institucional no período comparado."
            )

        add_paragrafo(doc, texto)

    add_paragrafo(
        doc,
        "A comparação longitudinal fortalece o uso do Teste de Progresso como instrumento de monitoramento "
        "contínuo da aprendizagem e como evidência de gestão acadêmica baseada em dados."
    )

    doc.save(caminho_word)

    return caminho_excel, caminho_grafico, caminho_word


# =====================================================
# MAIN
# =====================================================

def main():
    config = carregar_config()

    comparar, planilhas_comparacao = perguntar_comparacao_longitudinal()

    pastas = criar_pasta_projeto(config)

    caminho_planilha = pedir_planilha(config)

    shutil.copy(
        caminho_planilha,
        os.path.join(pastas["original"], os.path.basename(caminho_planilha))
    )

    base = carregar_dados(caminho_planilha)

    coluna_nome, coluna_periodo, coluna_total = detectar_colunas(base)

    base, coluna_nome, coluna_periodo = preparar_base(
        base,
        coluna_nome,
        coluna_periodo,
        coluna_total,
        config
    )

    (
        ranking_alunos,
        ranking_periodos,
        tabela_ciclos,
        alunos_risco,
        pontos_extras,
        plano_acao,
        caminho_excel,
        tabelas_areas
    ) = gerar_tabelas(
        base,
        coluna_nome,
        coluna_periodo,
        pastas,
        config
    )

    graficos = gerar_graficos(
        base,
        ranking_periodos,
        tabela_ciclos,
        coluna_periodo,
        pastas
    )

    relatorio = criar_relatorio_word(
        base,
        ranking_periodos,
        tabela_ciclos,
        alunos_risco,
        pontos_extras,
        plano_acao,
        graficos,
        coluna_nome,
        coluna_periodo,
        pastas,
        tabelas_areas,
        config
    )

    if comparar:
        caminho_excel_long, caminho_grafico_long, caminho_word_long = executar_comparacao_longitudinal(
            base,
            alunos_risco,
            pontos_extras,
            planilhas_comparacao,
            pastas,
            config
        )

        print("\n==============================================")
        print("COMPARAÇÃO LONGITUDINAL GERADA COM SUCESSO!")
        print("==============================================")
        print(f"Tabela: {caminho_excel_long}")
        print(f"Gráfico: {caminho_grafico_long}")
        print(f"Relatório: {caminho_word_long}")

    print("\n==============================================")
    print("ROBÔ MEC FINALIZADO COM SUCESSO!")
    print("==============================================")
    print(f"Pasta do projeto: {pastas['base']}")
    print(f"Tabelas: {caminho_excel}")
    print(f"Relatório: {relatorio}")
    print(f"Gráficos: {pastas['graficos']}")
    print("==============================================")


if __name__ == "__main__":
    main()
